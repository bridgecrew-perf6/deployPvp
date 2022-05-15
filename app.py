#import required  packages
from flask import Flask ,request
#import docx
import logging
import uuid
import os
import shutil
from azure.storage.blob import BlobServiceClient, BlobClient, ContainerClient, __version__
app = Flask(__name__)
import json
import subprocess
#import logging



#import methods
from predictiveAnalysis import textAndCharacteristics
from predictiveAnalysis import removeRedundantCharacterists
from predictiveAnalysis import newPdfGetTextLabelsAndTextIndex
from predictiveAnalysis import classGrouping
from predictiveAnalysis import clusterFormation
from predictiveAnalysis import removeOtherClassCharacteristics
from predictiveAnalysis import uniqueTitleCharacteristics
from docxAndPdfGenerator import previewPdf
from docxAndPdfGenerator import createDocxWithDeliminator
from docxAndPdfGenerator import individualAbstractDocxCreation
from docxAndPdfGenerator import convertDocxToPdfEachAbstract
from docxAndPdfGenerator import folderCreation
from grobid_client.grobid_client import GrobidClient




@app.route('/', methods=['GET'])
def hello_world():
    args = request.args
    containerName = args.get("name")
    print(containerName)
    os.mkdir(containerName)

    connect_str = "DefaultEndpointsProtocol=https;AccountName=pvpstorage1;AccountKey=tTqnRrNj2jLyxkP3sdYaGraaNWqofISf8SQb71RKC4sSnOmYk0rPcDm0xfvAwpffSqOZ9SVaJ+TO+AStQ1rZnA==;EndpointSuffix=core.windows.net"
    
    container = ContainerClient.from_connection_string(connect_str, container_name=containerName)
    blob_list = container.list_blobs()
    for blob in blob_list:
        blob_name= blob.name
        blob_client = container.get_blob_client(blob.name)
        with open("./"+ containerName + "/" + blob.name, "wb") as my_blob:
            blob_data = blob_client.download_blob()
            blob_data.readinto(my_blob)

     

    configurationFile = open('./config.json')
    configData = json.load(configurationFile)
    
    
    #Loading default values from the configuration file
    modelLocation = configData["modelLocation"]
    TokenizerLocation = configData["TokenizerLocation"]
    #fileLocation = configData["outputLocation"]


    
    try:
        #print("Succesfully opened the pdf and extracted the required details")
        logging.info("Succesfully opened the   "+ blob_name  + containerName+ "pdf and extracted the required details")
        text , characteristics, docObject = textAndCharacteristics("./"+containerName+"/"+blob_name, " ")
        
    except Exception as e:
        
        logging.error(" Failed to open the  " + blob_name +containerName+  "pdf and could not extract the required details. ") 
        #print(e)


    #remove the redundant Characteristics if present
    characteristics = removeRedundantCharacterists(characteristics)

    #get the Labels for each paragraph
    textLabels , textIndex= newPdfGetTextLabelsAndTextIndex(text,characteristics,modelLocation,TokenizerLocation)


    #group all of them to differenct class
    title, author, affliation , abstract , noise , ids = classGrouping(textIndex, textLabels, text, characteristics)
    
    #form clusters for each classes
    titleCluster = clusterFormation(title , "Title")
    authorCluster = clusterFormation(author , "Author")
    affliationCluster = clusterFormation(affliation, "Affliation")
    abstractCluster = clusterFormation(abstract, "Abstract")
    noiseCluster = clusterFormation(noise, "Noise")
    idCluster = clusterFormation(ids, "Ids")    

    otherCharacteristics = []
    otherCharacteristics = removeOtherClassCharacteristics(authorCluster,otherCharacteristics)
    otherCharacteristics = removeOtherClassCharacteristics(affliationCluster,otherCharacteristics)
    otherCharacteristics = removeOtherClassCharacteristics(abstractCluster,otherCharacteristics)
    print() 
    print() 
    #Extract title characteristics
    try:
            #print("Compeleted Cluster formation")
            logging.error("Completed creating clusters  for  " + blob_name +containerName+ "pdf. ")            
            fontSize, fontFamily, colour, titleCharacteristicsList = uniqueTitleCharacteristics(titleCluster,otherCharacteristics)  
            logging.error("Learned the characteristics of the   " + blob_name +containerName + "pdf. ")            
        
    except Exception as e:
            #pass
            logging.error(" Failed to learn the characteristics of the   " + blob_name +containerName + "pdf. ")            
            #print(e)
    
    #creates docx file with deliminator
    try:
            #print("Preview Docx file for pdf started created")
            createDocxWithDeliminator(text , characteristics,titleCharacteristicsList,blob_name,"./"+containerName+"/")
            logging.error(" Generated the Complete Docx with deliminator for the whole   " + blob_name +containerName  + "pdf. ")            
            #print()

    except Exception as e:
            #pass          
            logging.error("Failed To create Complete Docx with deliminator" + blob_name +containerName  + "pdf. ")            
            print(e)     
            
    
    #creates docx file for each abstract
    try:    

            countOfDocx = individualAbstractDocxCreation(text ,characteristics,titleCharacteristicsList,containerName ,"./",docObject)
            #print("Completed creating each abstract docx file")
            logging.error(" Completed creating each abstract   for  " + blob_name + containerName + "pdf. ")            
            
    except Exception as e:
    
            #pass
            logging.error(" Failed to  create  each abstract  for  " +  blob_name + containerName + "pdf. ")            
            print(e)
            
    
    

    try:
        for i in range(1, countOfDocx + 1):        
                subprocess.check_output(['libreoffice', '--convert-to', 'pdf' ,'--outdir','./' + containerName,'./' + containerName + '/' +str(i) +'.docx',])
    except Exception as e:
                logging.error("Could not convert docx to pdf")
        

    #if __name__ == "__main__":
    client = GrobidClient(config_path="./config2.json")
    client.process("processFulltextDocument", "./"+ containerName  )


    


        #upload  xml files to azure blob
    blob_service_client = BlobServiceClient.from_connection_string(connect_str)
    xmlFiles = os.listdir("./"+containerName)
    for xmlFile in xmlFiles:
        if ".xml"  in xmlFile:
                print(xmlFile)
                blob_client = blob_service_client.get_blob_client(container=containerName, blob=str(xmlFile))
                with open("./"+ containerName+ "/" + xmlFile, "rb") as data:
                        blob_client.upload_blob(data)







    

    '''
    #each docx files are converted to pdf 
    try:

            convertDocxToPdfEachAbstract(countOfDocx,containerName, "./")   
            logging.error(" Completed creating pdf   for  each abstract for " +blob_name + containerName + "pdf. ")            
            logging.error(blob_name + containerName + "pdf completed!! ")
            #print("Process completed")
    except Exception as e:

            #print(e)

            logging.error(" Failed to  create pdf   for  each abstract for " + blob_name + containerName + "pdf. ")            
            logging.error(blob_name + containerName +" pdf failed")
            #print("Failed to  convert docx to pdf") 
    
    
    '''




    '''

    doc = docx.Document()
    doc.add_paragraph('first item in ordered list', style='List Number')
    doc.save('./'+ containerName+'/demo.docx')
    '''


    '''        
    logging.info('Python HTTP trigger function processed a request.')
    # Create the BlobServiceClient object which will be used to create a container client
    connect_str = "DefaultEndpointsProtocol=https;AccountName=pvpstorage1;AccountKey=tTqnRrNj2jLyxkP3sdYaGraaNWqofISf8SQb71RKC4sSnOmYk0rPcDm0xfvAwpffSqOZ9SVaJ+TO+AStQ1rZnA==;EndpointSuffix=core.windows.net" 
    blob_service_client = BlobServiceClient.from_connection_string(connect_str)
    logging.info("connection Established")
    # Create a unique name for the container
    container_name = str(uuid.uuid4())
    #print(container_name)
            
    #doc_para = doc.add_paragraph('Your paragraph goes here, ')
    logging.info("Connection establised")
    
    # Create the container
    container_client = blob_service_client.create_container(container_name)
    blob_client = blob_service_client.get_blob_client(container=container_name, blob="demo.docx")


    with open("demo.docx", "rb") as data:
        blob_client.upload_blob(data)
    logging.info("uploaded the blob")

    '''
    #shutil.rmtree(containerName)
    return "<p>" + str(containerName) + "Hello, World!</p>"


















'''
import docx

a= 3
b= 5
print(a+b)
print("success")
doc = docx.Document()
doc.add_paragraph(
    'first item in ordered list pvp' ,
)

doc.save('demo.docx')

doc = docx.Document('demo.docx')
for docpara in doc.paragraphs:
    print(docpara.text)


'''













