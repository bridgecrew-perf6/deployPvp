import docx
from docx2pdf import convert
import os
import datetime




# prints the pdf with the deliminater
def previewPdf(text , characteristics,titleCharacteristicsList):
    keyWords =  ["Introduction:","Objectives","Methods","Objective","Results","Conclusion","Background:","Objectives:","Methods:","Results:","Conclusion:","DISCUSSION:","Conclusions"]
    for i in range(len(text)):
        for k in range(len(characteristics[i])):
            if(characteristics[i][k] == titleCharacteristicsList):
                texts = text[i].split()
                if(texts[0] not in keyWords):
                    print("************************************Title***************************************************************")
            
        print(text[i])
    


#create Docx file with deliminator
def createDocxWithDeliminator(text , characteristics,titleCharacteristicsList,filenameDocx,fileLocation):
    docPreview = docx.Document()
    keyWords =  ["Introduction:","Objectives","Methods","Objective","Results","Conclusion","Background:","Objectives:","Methods:","Results:","Conclusion:","DISCUSSION:","Conclusions"]
    indexs = []
    
    for i in range(len(text)):
        for k in range(len(characteristics[i])):
            if(characteristics[i][k] == titleCharacteristicsList):
                texts = text[i].split()
                if(texts[0] not in keyWords):
                    indexs.append(i)
                    #docPreview.add_paragraph("************************************Title***************************")
                    
        
        #docPreview.add_paragraph(text[i])
    
        
    
    
    length = len(indexs)
    sequence = []
    for i in range(1 , length):
        if(indexs[ length  - i ] == indexs[length - i -1] + 1):
            sequence.append(indexs[length - i])
            #print(indexs[length - i])
        
        else:
            if(len(sequence) != 0 and len(sequence) <=4):
                for k in sequence:
                    indexs.remove(k)
                sequence = []
            if(len(sequence) != 0 and len(sequence) >4):
                   sequence = []
        
    #print(indexs)
    
      
        
        
    for i in range(len(text)):
        if(i in indexs):
            docPreview.add_paragraph("************************************Title***************************")
            
        docPreview.add_paragraph(text[i])

        
    print(fileLocation+filenameDocx+"Preview"+".docx")
    docPreview.save(fileLocation+filenameDocx+"Preview"+".docx")




def individualAbstractDocxCreation(text ,characteristics,titleCharacteristicsList, folder,fileLocation,docObject):
    keyWords =  ["Introduction:","Objectives","Methods","Objective","Results","Conclusion","Background:","Objectives:","Methods:","Results:","Conclusion:","DISCUSSION:","Conclusions"]
    indexs = []
    for i in range(len(text)):
        for k in range(len(characteristics[i])):
            if(characteristics[i][k] == titleCharacteristicsList):
                texts = text[i].split()
                if(len(texts) != 0):
                    if(texts[0] not in keyWords):
                        indexs.append(i)
    
    
    
    
    #print(indexs)
    
    
    
    
    
    
    
    
    
    
    
    length = len(indexs)
    sequence = []
    for i in range(1 , length):
        if(indexs[ length  - i ] == indexs[length - i -1] + 1):
            sequence.append(indexs[length - i])
            #print(indexs[length - i])
        
        else:
            if(len(sequence) != 0 and len(sequence) <=4):
                for k in sequence:
                    indexs.remove(k)
                sequence = []
            if(len(sequence) != 0 and len(sequence) >4):
                   sequence = []
        
    #print(indexs)
    
    
    
    

    
    
    
    
    abstractLoc = []
    for i in range(len(indexs)-1):
        Loc =[]
        Loc.append(indexs[i])
        Loc.append(indexs[i+1] -1)
        abstractLoc.append(Loc)

    Loc = [indexs[len(indexs)-1] ,len(text)-1]
    abstractLoc.append(Loc)
    
    count=0
    paraCount =0
    pageNumber = 0
    for page in docObject:
        pageNumber += 1 
        #print("H")
        blocks = page.getText("dict")["blocks"]
        for b in blocks:
            #paragraph = ""
            #character = []
            if b['type'] == 0:  # block contains text
                #print(b)
                #print()

                if(paraCount == abstractLoc[count][0]): 
                    doc = docx.Document()
                         
                
                if(paraCount>= abstractLoc[count][0] and    paraCount<= abstractLoc[count][1]):  
                    doc_para = doc.add_paragraph()
                    for l in b["lines"]:  # iterate through the text lines
                        #print("line")
                        #print(l)
                        for s in l["spans"]:  
                            #print(s)
                            if(int(s["size"]) <=6):
                                super_text = doc_para.add_run(s["text"]+"  ")
                                super_text.font.superscript = True
                                
                                #print(s)
                            else:
                                doc_para.add_run("  "+s["text"])
                                
                    
                    
                    
                    
                if(paraCount == abstractLoc[count][1]):
                    doc.save(fileLocation+folder+"/"+str(count+1)+".docx")
                    count +=1
                    
                paraCount +=1 
       
    return count
    #print(count)                

    
def convertDocxToPdfEachAbstract(countOfDocx,folder,fileLocation):
    for i in range(1,countOfDocx + 1):
        print(fileLocation+folder+"/"+str(i)+".docx" + "    is geting converted to pdf Please wait")
        convert(fileLocation+folder+"/"+str(i)+".docx")
    print("Succesfully completed all the docx to pdf")
    
    
def folderCreation(fileName, fileLocation):
    end = len(fileName) - 4
    start =0
    for i in range(len(fileName)):
        if(fileName[end - 1 - i] == '/'):
            start = end -i
            break
    folderName = fileName[start: end]
    currentTime = datetime.datetime.now()
    currentTime = str(currentTime)
    time = ""
    for i in currentTime:
        if(i == '.'):
            pass

        elif(i == ' '):
            time +='-'
        elif(i==':'):
            time += '-'
            
        else:
            time += i

    #print(time)
        
    
    
    folderName = time + folderName
    
    
    os.chdir(fileLocation)
    os.mkdir(folderName)
    return folderName

